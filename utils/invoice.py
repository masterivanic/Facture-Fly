import os
import subprocess
from io import BytesIO
from pathlib import Path

import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from PIL import Image

REPORTS_DIR = Path(__file__).resolve().parent.parent / "generated_invoice"


class InvoiceGenerator:
    @staticmethod
    def generate_pretty_invoice(
        data,
        logo_url: str = "https://example.com/path/to/logo.png",
        footer_message: str = "Thank you for your business! You're awesome!",
    ) -> BytesIO:
        doc = Document()

        section = doc.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        header = section.header
        header_paragraph = header.paragraphs[0]

        try:
            response = requests.get(logo_url)
            if response.status_code == 200:
                logo_stream = BytesIO(response.content)
                logo_image = Image.open(logo_stream)
                logo_width, logo_height = logo_image.size

                # Resize the logo to fit the header
                max_width = Inches(2)  # Adjust as needed
                if logo_width > max_width:
                    ratio = max_width / logo_width
                    logo_width = max_width
                    logo_height = logo_height * ratio

                header_paragraph.add_run().add_picture(logo_stream, width=logo_width)
        except Exception as e:
            print(f"Failed to fetch logo: {e}")

        title = doc.add_paragraph()
        title_run = title.add_run("INVOICE")
        title_run.font.name = "Arial"
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        title_run.bold = True
        title.alignment = 1

        doc.add_paragraph().add_run().add_break()

        doc.add_paragraph(f"Invoice Number: {data['label']}")
        doc.add_paragraph(f"Date: {data['emission_date']}")
        doc.add_paragraph(f"Due Date: {data['due_date']}")
        doc.add_paragraph(f"Customer: {data['customer_name']}")

        doc.add_paragraph().add_run().add_break()
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Item"
        hdr_cells[1].text = "Quantity"
        hdr_cells[2].text = "Unit Price"
        hdr_cells[3].text = "Total"

        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "4F81BD")
            cell._tc.get_or_add_tcPr().append(shading)

        for article in data["articles"]:
            row_cells = table.add_row().cells
            row_cells[0].text = article["article_label"]
            row_cells[1].text = str(article["article_quantity"])
            row_cells[2].text = f"${article['article_price']:.2f}"
            row_cells[
                3
            ].text = f"${article['article_quantity'] * article['article_price']:.2f}"

        doc.add_paragraph().add_run().add_break()
        doc.add_paragraph(f"Subtotal: ${data['amount']:.2f}")
        doc.add_paragraph(f"Discount: ${data['discount']:.2f}")
        doc.add_paragraph(f"Tax: ${data['taxe']:.2f}")
        doc.add_paragraph(f"Total: ${data['paid_amount']:.2f}")

        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = footer_message
        footer_paragraph.alignment = 1
        footer_paragraph.runs[0].font.italic = True

        docx_file = os.path.join(REPORTS_DIR, "temp_invoice.docx")
        doc.save(docx_file)

        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                REPORTS_DIR,
                docx_file,
            ]
        )
        pdf_file = os.path.join(REPORTS_DIR, "temp_invoice.pdf")
        with open(pdf_file, "rb") as f:
            pdf_stream = BytesIO(f.read())
        return pdf_stream
